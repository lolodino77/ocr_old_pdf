from docx import Document
import os

def check_if_modernized_txt_already_exported_to_word(document_name, directory='.'):
    """
    Vérifie si le texte modernisé existe déjà pour un document donné.
    """
    # Retirer l'extension .docx si présente
    base_name = os.path.splitext(document_name)[0]

    # Construire le nom du fichier modernisé attendu
    # Par exemple, si le document est "Daille_S_133_Noel", 
    # on renomme en "Daille_S_133_Noel_modernized_cleaned_text.docx"
    if("_modernized_cleaned_text" not in document_name):
        document_name = f"{base_name}_modernized_cleaned_text.docx"

    # Construire le chemin absolu complet
    full_path = os.path.abspath(os.path.join(directory, document_name))

    # Debug utile
    print(f"Répertoire courant réel : {os.getcwd()}")
    print(f"Recherche du fichier : {full_path}")
    # print(f"Fichiers présents dans {os.path.abspath(directory)} : {os.listdir(os.path.abspath(directory))}")

    # Vérification d'existence
    return os.path.exists(full_path)

def convert_modernized_txt_to_word(input_txt_file, output_docx_file):
    doc = Document()

    # Lire le fichier texte complet
    with open(input_txt_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Regrouper les paragraphes (séparés par lignes vides)
    paragraph = []
    for line in lines:
        stripped = line.strip()
        if stripped:
            paragraph.append(stripped)
        else:
            # Ligne vide = fin de paragraphe
            if paragraph:
                doc.add_paragraph(" ".join(paragraph))
                paragraph = []

    # Ajouter le dernier paragraphe s’il existe
    if paragraph:
        doc.add_paragraph(" ".join(paragraph))

    # Sauvegarde du fichier texte désormais sans ligne vide entre chaque paragraphe
    temp_txt_file = "temp_no_double_spaces.txt"
    with open(temp_txt_file, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")

    # Sauvegarde du document Word
    doc.save(output_docx_file)
    print(f"✅ Le fichier '{input_txt_file}' a été converti en '{output_docx_file}' sans double espaces.")
